from docx import Document
from docx.shared import Inches

document = Document()

# Profile Picture:

document.add_picture(
    'foto.png',
    width=Inches(2.0)
)

# Information about name, phone number and email

name = input('What is your name? ')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# About me...

document.add_heading('About me')
document.add_paragraph(
    input('Tell me about yourself! ')
)

# Work experience

document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter the name of the company: ')
from_date = input('From date: ')
to_date = input('To_date ')

p.add_run(company + ' ').bold = True
p.add_run('(' + from_date + ' - ' + to_date + ')\n').italic = True

experience_details = input(
    'Describe your experience at ' + company + ': '
)
p.add_run(experience_details)

#Add more experiences

while True:
    has_more_experiences = input(
        'Do you have more experiences? Y or N \n')
    if has_more_experiences.upper() == 'Y':
        p = document.add_paragraph()

        company = input('Enter the name of the company: ')
        from_date = input('From date: ')
        to_date = input('To_date ')

        p.add_run(company + ' ').bold = True
        p.add_run('(' + from_date + ' - ' + to_date + ')\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ': \n'
        )
        p.add_run(experience_details)
    else:
        break

#Add Skills
document.add_heading('Skills')
skill = input('Enter a Skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

#Add more skills
while True:
    has_more_skills = input('Do you have more Skills? Y or N: \n')
    if (has_more_skills.upper() == 'Y'):
        skill = input('Enter Skill: \n')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated from Python Script as a coding practice'

document.save('cv.docx')